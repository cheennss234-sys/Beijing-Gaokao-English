#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成 2024-2026 北京高三英语试卷课外高频词汇统计 Word 文档。
应用 Anthropic 官方品牌配色与字体（brand-guidelines）。
"""

import json
import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============ Anthropic 品牌配色 ============
COLOR_DARK = RGBColor(0x14, 0x14, 0x13)          # 主文本与深色背景
COLOR_LIGHT = RGBColor(0xFA, 0xF9, 0xF5)         # 浅色背景与深色上的文本
COLOR_MID_GRAY = RGBColor(0xB0, 0xAE, 0xA5)      # 次要元素
COLOR_LIGHT_GRAY = RGBColor(0xE8, 0xE6, 0xDC)    # 微妙背景
COLOR_ORANGE = RGBColor(0xD9, 0x77, 0x57)        # 主强调色
COLOR_BLUE = RGBColor(0x6A, 0x9B, 0xCC)          # 次强调色
COLOR_GREEN = RGBColor(0x78, 0x8C, 0x5D)         # 三级强调色

# 16 进制字符串（用于单元格底纹 shading）
HEX_DARK = "141413"
HEX_LIGHT = "FAF9F5"
HEX_MID_GRAY = "B0AEA5"
HEX_LIGHT_GRAY = "E8E6DC"
HEX_ORANGE = "D97757"
HEX_BLUE = "6A9BCC"
HEX_GREEN = "788C5D"

# ============ Anthropic 品牌字体 ============
FONT_HEADING = "Poppins"      # 标题（回退 Arial）
FONT_BODY = "Lora"            # 正文（回退 Georgia）
FONT_HEADING_FALLBACK = "Arial"
FONT_BODY_FALLBACK = "Georgia"


# ============ 工具函数 ============
def set_cell_shading(cell, hex_color):
    """设置单元格背景色（shading）。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_paragraph_shading(paragraph, hex_color):
    """设置段落背景色（shading）。"""
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    p_pr.append(shd)


def set_run_font(run, font_name, fallback_font, size_pt, color=None, bold=False):
    """设置 run 的字体（含东亚字体）、字号、颜色、加粗。"""
    run.font.name = font_name
    # 设置东亚字体（中文显示）
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn('w:rFonts'))
    if r_fonts is None:
        r_fonts = OxmlElement('w:rFonts')
        r_pr.append(r_fonts)
    r_fonts.set(qn('w:ascii'), font_name)
    r_fonts.set(qn('w:hAnsi'), font_name)
    r_fonts.set(qn('w:eastAsia'), fallback_font)
    r_fonts.set(qn('w:cs'), font_name)
    run.font.size = Pt(size_pt)
    if color is not None:
        run.font.color.rgb = color
    run.font.bold = bold


def set_paragraph_spacing(paragraph, before_pt=0, after_pt=0, line_pt=None):
    """设置段落间距。"""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    if line_pt is not None:
        pf.line_spacing = line_pt


def add_paragraph_border(paragraph, hex_color=HEX_LIGHT_GRAY, size=6):
    """为段落添加边框（用于说明区块的视觉分隔）。"""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    for edge in ('top', 'left', 'bottom', 'right'):
        bdr = OxmlElement(f'w:{edge}')
        bdr.set(qn('w:val'), 'single')
        bdr.set(qn('w:sz'), str(size))
        bdr.set(qn('w:space'), '4')
        bdr.set(qn('w:color'), hex_color)
        p_bdr.append(bdr)
    p_pr.append(p_bdr)


# ============ 主流程 ============
def main():
    input_path = "/workspace/output/vocab_data.json"
    output_path = "/workspace/output/beijing_grade12_english_vocab.docx"

    # 1. 读取 JSON 数据
    with open(input_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    meta = data["metadata"]
    vocab_list = data["vocab_list"]
    total_words = meta["total_words"]

    # 2. 创建文档
    doc = Document()

    # 设置默认正文样式（Lora 回退 Georgia）
    normal_style = doc.styles['Normal']
    normal_style.font.name = FONT_BODY
    normal_style.font.size = Pt(11)
    r_pr = normal_style.element.get_or_add_rPr()
    r_fonts = r_pr.find(qn('w:rFonts'))
    if r_fonts is None:
        r_fonts = OxmlElement('w:rFonts')
        r_pr.append(r_fonts)
    r_fonts.set(qn('w:ascii'), FONT_BODY)
    r_fonts.set(qn('w:hAnsi'), FONT_BODY)
    r_fonts.set(qn('w:eastAsia'), FONT_BODY_FALLBACK)
    r_fonts.set(qn('w:cs'), FONT_BODY)

    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # ============ 3. 标题 ============
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title_para, before_pt=6, after_pt=6, line_pt=1.5)
    title_run = title_para.add_run(meta["description"])
    set_run_font(title_run, FONT_HEADING, FONT_HEADING_FALLBACK,
                 size_pt=24, color=COLOR_DARK, bold=True)

    # 标题下方副标题（年份范围 + 词条总数）
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(subtitle_para, before_pt=0, after_pt=12, line_pt=1.2)
    subtitle_run = subtitle_para.add_run(
        f"年份范围：{meta['years']}    |    词条总数：{total_words} 个"
    )
    set_run_font(subtitle_run, FONT_BODY, FONT_BODY_FALLBACK,
                 size_pt=11, color=COLOR_MID_GRAY, bold=False)

    # ============ 4. 说明部分（Light Gray 背景段落） ============
    # 说明区块标题
    desc_title_para = doc.add_paragraph()
    set_paragraph_spacing(desc_title_para, before_pt=6, after_pt=4, line_pt=1.2)
    set_paragraph_shading(desc_title_para, HEX_LIGHT_GRAY)
    add_paragraph_border(desc_title_para, hex_color=HEX_LIGHT_GRAY, size=4)
    desc_title_run = desc_title_para.add_run("数据说明")
    set_run_font(desc_title_run, FONT_HEADING, FONT_HEADING_FALLBACK,
                 size_pt=14, color=COLOR_DARK, bold=True)

    # 说明内容段落
    regions_str = "、".join(meta["regions"])
    exam_types_str = "、".join(meta["exam_types"])
    info_lines = [
        ("年份范围", meta["years"]),
        ("覆盖城区", regions_str),
        ("考试类型", exam_types_str),
        ("课本基准", meta["textbook_baseline"]),
        ("最低词频", f"{meta['min_frequency']} 次"),
        ("词条总数", f"{total_words} 个"),
    ]

    for label, value in info_lines:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before_pt=0, after_pt=2, line_pt=1.3)
        set_paragraph_shading(p, HEX_LIGHT_GRAY)
        # 标签
        r1 = p.add_run(f"【{label}】")
        set_run_font(r1, FONT_BODY, FONT_BODY_FALLBACK,
                     size_pt=10.5, color=COLOR_ORANGE, bold=True)
        # 值
        r2 = p.add_run(f"  {value}")
        set_run_font(r2, FONT_BODY, FONT_BODY_FALLBACK,
                     size_pt=10.5, color=COLOR_DARK, bold=False)

    # 数据覆盖说明（coverage_notes）
    notes_label_para = doc.add_paragraph()
    set_paragraph_spacing(notes_label_para, before_pt=6, after_pt=2, line_pt=1.3)
    set_paragraph_shading(notes_label_para, HEX_LIGHT_GRAY)
    r_lbl = notes_label_para.add_run("【数据覆盖说明】")
    set_run_font(r_lbl, FONT_BODY, FONT_BODY_FALLBACK,
                 size_pt=10.5, color=COLOR_ORANGE, bold=True)

    notes_para = doc.add_paragraph()
    set_paragraph_spacing(notes_para, before_pt=0, after_pt=10, line_pt=1.5)
    set_paragraph_shading(notes_para, HEX_LIGHT_GRAY)
    add_paragraph_border(notes_para, hex_color=HEX_LIGHT_GRAY, size=4)
    r_notes = notes_para.add_run(meta["coverage_notes"])
    set_run_font(r_notes, FONT_BODY, FONT_BODY_FALLBACK,
                 size_pt=10, color=COLOR_DARK, bold=False)

    # ============ 5. 词表表格 ============
    # 表格标题
    table_title_para = doc.add_paragraph()
    set_paragraph_spacing(table_title_para, before_pt=10, after_pt=6, line_pt=1.2)
    table_title_run = table_title_para.add_run("高频词汇表")
    set_run_font(table_title_run, FONT_HEADING, FONT_HEADING_FALLBACK,
                 size_pt=16, color=COLOR_DARK, bold=True)

    # 创建表格（表头 + 数据行）
    headers = ["排名", "单词", "词性", "中文释义", "频次"]
    table = doc.add_table(rows=1 + len(vocab_list), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # 设置列宽
    col_widths = [Cm(1.5), Cm(3.5), Cm(1.8), Cm(7.0), Cm(1.8)]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = col_widths[idx]

    # 表头行
    header_row = table.rows[0]
    for idx, header_text in enumerate(headers):
        cell = header_row.cells[idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, HEX_ORANGE)
        # 清空默认段落并写入
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before_pt=2, after_pt=2, line_pt=1.2)
        r = p.add_run(header_text)
        set_run_font(r, FONT_HEADING, FONT_HEADING_FALLBACK,
                     size_pt=11, color=COLOR_LIGHT, bold=True)

    # 数据行（斑马纹）
    for i, item in enumerate(vocab_list):
        row = table.rows[i + 1]
        # 交替背景色
        bg_hex = HEX_LIGHT if i % 2 == 0 else HEX_LIGHT_GRAY

        values = [
            str(item["rank"]),
            item["word"],
            item["pos"],
            item["meaning"],
            str(item["frequency"]),
        ]

        for idx, val in enumerate(values):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_shading(cell, bg_hex)
            cell.text = ""
            p = cell.paragraphs[0]
            # 排名、词性、频次居中；单词、释义左对齐
            if idx in (0, 2, 4):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(p, before_pt=1, after_pt=1, line_pt=1.2)
            r = p.add_run(val)

            # 颜色规则：单词列 Dark，频次列 Orange 强调，其余 Dark
            if idx == 1:  # 单词列
                set_run_font(r, FONT_BODY, FONT_BODY_FALLBACK,
                             size_pt=11, color=COLOR_DARK, bold=True)
            elif idx == 4:  # 频次列
                set_run_font(r, FONT_HEADING, FONT_HEADING_FALLBACK,
                             size_pt=11, color=COLOR_ORANGE, bold=True)
            elif idx == 0:  # 排名列
                set_run_font(r, FONT_BODY, FONT_BODY_FALLBACK,
                             size_pt=10.5, color=COLOR_MID_GRAY, bold=False)
            else:  # 词性、释义
                set_run_font(r, FONT_BODY, FONT_BODY_FALLBACK,
                             size_pt=10.5, color=COLOR_DARK, bold=False)

    # 设置表格边框
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        bdr = OxmlElement(f'w:{edge}')
        bdr.set(qn('w:val'), 'single')
        bdr.set(qn('w:sz'), '4')
        bdr.set(qn('w:space'), '0')
        bdr.set(qn('w:color'), HEX_LIGHT_GRAY)
        tbl_borders.append(bdr)
    tbl_pr.append(tbl_borders)

    # ============ 6. 页脚 ============
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gen_date = datetime.now().strftime("%Y-%m-%d")
    footer_text = f"共 {total_words} 个词条 | 生成日期：{gen_date}"
    footer_run = footer_para.add_run(footer_text)
    set_run_font(footer_run, FONT_BODY, FONT_BODY_FALLBACK,
                 size_pt=9, color=COLOR_MID_GRAY, bold=False)

    # 7. 保存文档
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

    # 返回统计信息
    file_size = os.path.getsize(output_path)
    print(f"生成路径: {output_path}")
    print(f"词条总数: {total_words}")
    print(f"文件大小: {file_size} 字节 ({file_size / 1024:.2f} KB)")


if __name__ == "__main__":
    main()
